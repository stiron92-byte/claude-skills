#!/usr/bin/env python3
"""DOCX 문서 텍스트 교정 — 서식/이미지 보존, 텍스트만 수정."""

from docx import Document
from copy import deepcopy

SRC = "/Users/gangdasol/Desktop/claude-code-skills-chapter2-doc-automation.docx"
DST = "/Users/gangdasol/Desktop/claude-code-skills-chapter2-doc-automation_수정본.docx"

# ──────────────────────────────────────────────
# 치환 목록: (원본, 수정)
# ──────────────────────────────────────────────
REPLACEMENTS = [
    # === 오탈자 ===
    ("외한 및 스왑", "외환 및 스왑"),
    ("멀 잘 못눌렀나", "뭘 잘못 눌렀나"),
    ("분들으 폴더에서", "분들은 폴더에서"),
    ("빼앗아 가는", "빼앗기는"),

    # === 주술 부조화 ===
    (
        "이 반복 패턴이 Cowork로 자동화해야 합니다.",
        "이 반복 작업을 Cowork로 자동화할 때입니다.",
    ),
    (
        "막막하게 느껴지는 것은 당연한 반응입니다.",
        "막막한 것은 자연스러운 일입니다.",
    ),

    # === AI투 줄이기: "~것입니다" 5연속 → 분산 ===
    (
        "첫 번째는 PDF 파일을 전달해 발표용 PPT를 만드는 것, 두 번째는 웹 주소(URL)를 전달해 슬라이드로 정리하는 것, 세 번째는 기업 정보가 담긴 파일을 바탕으로 정부지원사업 발표자료를 만드는 것, 네 번째는 내가 가진 양식에 맞춰 문서를 생성하는 것입니다. 네 가지 모두 Cowork 입력창에 말로 요청하는 것이 전부입니다.",
        "첫 번째는 PDF를 발표용 PPT로 변환하고, 두 번째는 웹 주소(URL)를 슬라이드로 정리하고, 세 번째는 기업 정보 파일들을 바탕으로 정부지원사업 발표자료를 구성하고, 네 번째는 내가 가진 양식에 맞춰 문서를 생성합니다. 네 가지 모두 Cowork 입력창에 말로 요청하면 됩니다.",
    ),

    # === AI투 줄이기: "핵심입니다" 3회 → 1회만 남기기 ===
    (
        "그 부분이 전체 소요 시간의 70~80%를 차지한다는 것이 문제의 핵심입니다.",
        "그 부분이 전체 소요 시간의 70~80%를 차지합니다.",
    ),
    (
        "이것은 이 장 전체를 관통하는 핵심입니다.",
        "이 장의 모든 실습이 보여주는 메시지이기도 합니다.",
    ),

    # === AI투 줄이기: "~다는 것을 / ~라는 것입니다" 패턴 ===
    (
        "커스터마이징까지 할 수 있다는 것을 이 책을 통해 배우시게 될 겁니다.",
        "커스터마이징까지 할 수 있다는 걸 이 책에서 직접 경험하게 됩니다.",
    ),
    (
        "클로드에게 맥락을 주는 것, 그것이 좋은 결과물을 만드는 가장 빠른 방법입니다.",
        "클로드에게 맥락을 주는 것이 좋은 결과물을 만드는 가장 빠른 방법입니다.",
    ),

    # === AI투 줄이기: "차이를 만드는 것은 도구가 아니라" ===
    (
        "차이를 만드는 것은 도구가 아니라 여러분이 클로드에게 얼마나 많은 맥락과 목적을 전달했느냐입니다.",
        "차이는 도구에서 나오지 않습니다. 클로드에게 얼마나 구체적인 맥락과 목적을 전달했느냐에서 나옵니다.",
    ),

    # === 어색한 문장 다듬기 ===
    (
        "일단, 위 네 가지 실습을 통해 그 의심을 직접 눈으로 풀어드리겠습니다.",
        "일단 네 가지 실습을 따라하면서 직접 확인해보시기 바랍니다.",
    ),
    (
        "마치 긴 회의록을 요약본으로 만들어두는 것처럼요.",
        "긴 회의록에서 핵심만 추려 요약본을 만드는 것과 비슷합니다.",
    ),
]


def replace_in_paragraph(paragraph, old_text, new_text):
    """단락의 run을 순회하며 텍스트를 치환. 서식(bold, italic, font 등)은 보존."""
    # 전체 텍스트에 대상이 없으면 스킵
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # 방법 1: 단일 run에 포함된 경우 (가장 안전)
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # 방법 2: 여러 run에 걸쳐있는 경우
    # run 텍스트를 이어붙여서 위치를 찾고, 해당 run들만 수정
    runs = paragraph.runs
    if not runs:
        return False

    # 각 run의 시작/끝 위치 계산
    positions = []
    pos = 0
    for run in runs:
        start = pos
        end = pos + len(run.text)
        positions.append((start, end))
        pos = end

    # old_text의 위치 찾기
    idx = full_text.find(old_text)
    if idx == -1:
        return False

    old_end = idx + len(old_text)

    # 영향받는 run 범위 찾기
    first_run = None
    last_run = None
    for i, (start, end) in enumerate(positions):
        if start <= idx < end and first_run is None:
            first_run = i
        if start < old_end <= end:
            last_run = i
            break

    if first_run is None or last_run is None:
        return False

    # 첫 번째 run에 교체 텍스트 넣기
    first_start = positions[first_run][0]
    last_end = positions[last_run][1]

    # 첫 run: old_text 시작 전 텍스트 + new_text + 마지막 run의 old_text 이후 텍스트
    prefix = runs[first_run].text[:idx - first_start]
    suffix = runs[last_run].text[old_end - positions[last_run][0]:]
    runs[first_run].text = prefix + new_text + suffix

    # 중간/마지막 run의 텍스트 비우기 (run 자체는 삭제하지 않음 — 서식 보존)
    for i in range(first_run + 1, last_run + 1):
        runs[i].text = ""

    return True


def has_image(paragraph):
    """단락에 이미지가 포함되어 있는지 확인."""
    # inline 이미지 확인
    if paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        return True
    if paragraph._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
        return True
    # run 내부의 drawing 확인
    for run in paragraph.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            return True
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            return True
    return False


def main():
    doc = Document(SRC)
    changes = []

    for old_text, new_text in REPLACEMENTS:
        found = False
        for para in doc.paragraphs:
            # 이미지가 있는 단락은 건드리지 않음
            if has_image(para):
                continue
            if old_text in para.text:
                result = replace_in_paragraph(para, old_text, new_text)
                if result:
                    found = True
                    changes.append(f"  ✅ '{old_text[:30]}...' → '{new_text[:30]}...'")

        # 표 안의 텍스트도 확인
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            result = replace_in_paragraph(para, old_text, new_text)
                            if result:
                                found = True
                                changes.append(f"  ✅ [표] '{old_text[:30]}...' → '{new_text[:30]}...'")

        if not found:
            changes.append(f"  ⚠️  미발견: '{old_text[:40]}...'")

    doc.save(DST)

    print(f"원본: {SRC}")
    print(f"수정본: {DST}")
    print(f"\n변경 사항 ({len(changes)}건):")
    for c in changes:
        print(c)


if __name__ == "__main__":
    main()
