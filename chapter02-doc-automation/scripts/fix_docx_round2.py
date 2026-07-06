#!/usr/bin/env python3
"""DOCX 문서 2차 교정 — 오탈자, 톤 통일, AI투 제거, 문맥 보완."""

from docx import Document

SRC = "/Users/gangdasol/Desktop/claude-code-skills-chapter2-doc-automation_수정본.docx"
DST = "/Users/gangdasol/Desktop/claude-code-skills-chapter2-doc-automation_수정본.docx"

# ──────────────────────────────────────────────
# 단순 치환 (run 내에서 str.replace)
# ──────────────────────────────────────────────
REPLACEMENTS = [
    # === A. 오탈자 ===
    # A1: 관광서 → 관공서
    ("관광서에서 사용한", "관공서에서 사용한"),
    # A2: "대해" → "에 대해"
    ("무엇인지 대해 감을", "무엇인지에 대해 감을"),
    # A3: 띄어쓰기
    ("잡으셨을겁니다", "잡으셨을 겁니다"),
    # A5: 구어체 정리
    ("사용할 수 있는거 이제 조금 감이 오시죠?", "사용할 수 있다는 거, 이제 조금 감이 오시죠?"),

    # === C. 톤/경어 통일 ===
    # C1: 과도한 높임
    ("아래 그림처럼 보이십니다", "아래 그림처럼 보입니다"),
    # C2: 반말형 나열
    ("chapter1에서 받았고 압축을 풀었던", "chapter1에서 받아서 압축을 풀어둔"),
    # C3: 띄어쓰기
    ("커피한잔 하고 오시면", "커피 한 잔 하고 오시면"),
    ("커피한잔을 하고 오시면", "커피 한 잔을 하고 오시면"),

    # === D. AI투 제거 ===
    # D1: 3중 명사절
    (
        "클로드를 잘 쓴다는 것은 복잡한 기술을 익히는 것이 아니라, 내가 원하는 것을 명확하게 전달하는 능력입니다.",
        "클로드를 잘 쓰는 방법은 간단합니다. 내가 원하는 것을 명확하게 전달하면 됩니다.",
    ),
    # D2: 3번째 반복 메시지 축약
    (
        "클로드에게 넘겨주는 자료가 많아질수록, 목적이 구체적일수록 결과물은 달라집니다. 이 장의 모든 실습이 보여주는 메시지이기도 합니다.",
        "앞선 실습에서도 확인했듯이, 넘겨주는 자료가 많고 목적이 구체적일수록 결과물의 깊이가 달라집니다.",
    ),
    # D3: AI 공감형 도입
    (
        "이 경험이 낯설지 않으신가요? 업무의 진짜 가치는 '무슨 내용을 전달할 것인가'를 고민하는 데 있는데, 실제 시간의 대부분은 '어떻게 예쁘게 포맷을 맞출 것인가'에 쏟아붓게 됩니다.",
        "익숙한 상황이실 겁니다. 정작 중요한 건 '무슨 내용을 전달할지'인데, 실제 시간은 대부분 '어떻게 예쁘게 포맷을 맞출지'에 쏟아붓게 됩니다.",
    ),

    # === B2: 줄바꿈 뜬금없는 언급 → 자연스럽게 ===
    (
        "혹시 줄바꿈은 어떻게 해요? 라는 질문을 하실 수 있는 shift+enter를 누르시면 됩니다.",
        "참고로 입력창에서 줄바꿈을 하려면 Shift+Enter를 누르면 됩니다.",
    ),

    # === B1: 직군별 시나리오 문맥 보완 ===
    # "살펴보겠습니다." 뒤에 예시 추가, "공통점이 보이시나요?" 앞에 연결
    (
        "여러분의 업무 상황과 겹치는 부분이 있는지 확인해보세요.",
        "여러분의 업무 상황과 겹치는 부분이 있는지 확인해보세요. 금융권은 시황 보고서를, 마케팅은 캠페인 결과를, 인사팀은 채용 현황을, 영업은 분기 실적을 PPT로 정리하는 일이 반복됩니다.",
    ),

    # === B3: 비교 실습 전환 문장 추가 ===
    (
        "지금까지 우리는 무역회사 입장에서 과장님께 보고드리는 상황을 가정하고 실습을 진행했습니다. 그런데 같은 PDF를 아무런 조건 없이",
        "지금까지 우리는 무역회사 입장에서 과장님께 보고드리는 상황을 가정하고 실습을 진행했습니다. 여기서 한 가지 더 살펴보겠습니다. 같은 PDF를 아무런 조건 없이",
    ),

    # === B4: 압축 설명 전환 문장 추가 ===
    (
        "Cowork에서 작업을 오래 하다 보면 가끔 이런 화면을 만나게 됩니다.",
        "실습을 마치기 전에 하나만 더 알아두면 좋은 것이 있습니다. Cowork에서 작업을 오래 하다 보면 가끔 이런 화면을 만나게 됩니다.",
    ),

    # === D4: "여러분" 과다 사용 — 일부를 자연스럽게 대체 ===
    # 선별적으로만 교체 (문맥에 맞는 것만)
    ("여러분의 가장 희소한 자원인 시간이 거기에 묶여 있습니다.", "가장 희소한 자원인 시간이 거기에 묶여 있는 셈입니다."),
    ("여러분이 이 파일들을 직접 열거나 수정할 필요는 전혀 없습니다.", "이 파일들을 직접 열거나 수정할 필요는 전혀 없습니다."),
    ("여러분이 직접 비교해볼 수 있도록", "직접 비교해볼 수 있도록"),
    ("여러분이 직접 경험하신 것이 바로 그것입니다.", "방금 직접 경험하신 것이 바로 그것입니다."),
    ("여러분의 상황에 맞게 조정하여", "본인의 상황에 맞게 조정하여"),
    ("여러분이 알려준 만큼 이해합니다.", "알려준 만큼 이해합니다."),
    ("여러분이 클로드에게 얼마나 많은 맥락과 목적을 전달했느냐입니다.", "클로드에게 얼마나 구체적인 맥락과 목적을 전달했느냐에서 나옵니다."),
    ("여러분은 결과물을 확인하고 수정을 요청하면 됩니다.", "결과물을 확인하고 수정을 요청하면 됩니다."),
    ("여러분이 직접 사용할 필요는 없습니다.", "직접 사용할 필요는 없습니다."),
    ("여러분들이 흔히 보실 수 있는", "흔히 볼 수 있는"),
]


def replace_in_paragraph(paragraph, old_text, new_text):
    """단락의 run에서 텍스트를 치환. 서식 보존."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # 단일 run에 포함된 경우
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # 여러 run에 걸친 경우
    runs = paragraph.runs
    if not runs:
        return False

    positions = []
    pos = 0
    for run in runs:
        start = pos
        end = pos + len(run.text)
        positions.append((start, end))
        pos = end

    idx = full_text.find(old_text)
    if idx == -1:
        return False

    old_end = idx + len(old_text)

    first_run = None
    last_run = None
    for i, (start, end) in enumerate(positions):
        if start <= idx < end and first_run is None:
            first_run = i
        if start < old_end <= end:
            last_run = i
            break

    if first_run is None or last_run is None:
        return False

    first_start = positions[first_run][0]
    prefix = runs[first_run].text[:idx - first_start]
    suffix = runs[last_run].text[old_end - positions[last_run][0]:]
    runs[first_run].text = prefix + new_text + suffix

    for i in range(first_run + 1, last_run + 1):
        runs[i].text = ""

    return True


def has_image(paragraph):
    """단락에 이미지가 포함되어 있는지 확인."""
    for ns in [
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict',
    ]:
        if paragraph._element.findall(f'.//{ns}'):
            return True
    for run in paragraph.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            return True
    return False


def main():
    doc = Document(SRC)
    changes = []

    for old_text, new_text in REPLACEMENTS:
        found = False
        for para in doc.paragraphs:
            if has_image(para):
                continue
            if old_text in para.text:
                result = replace_in_paragraph(para, old_text, new_text)
                if result:
                    found = True
                    label = old_text[:35].replace('\n', ' ')
                    changes.append(f"  ✅ '{label}...'")

        # 표 안도 확인
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            result = replace_in_paragraph(para, old_text, new_text)
                            if result:
                                found = True
                                changes.append(f"  ✅ [표] '{old_text[:35]}...'")

        if not found:
            changes.append(f"  ⚠️  미발견: '{old_text[:40]}...'")

    doc.save(DST)

    print(f"변경 사항 ({len(changes)}건):")
    for c in changes:
        print(c)

    # 검증
    doc2 = Document(DST)
    def count_images(d):
        return sum(1 for r in d.part.rels.values() if 'image' in r.reltype)

    print(f"\n검증:")
    print(f"  단락: {len(doc2.paragraphs)}  ✅")
    print(f"  이미지: {count_images(doc2)}  ✅")
    print(f"  표: {len(doc2.tables)}  ✅")


if __name__ == "__main__":
    main()
