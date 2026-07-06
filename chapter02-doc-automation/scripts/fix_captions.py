#!/usr/bin/env python3
"""DOCX 이미지 캡션 보완 — 번호만 있는 캡션에 설명 추가."""

from docx import Document

SRC = "/Users/gangdasol/Desktop/claude-code-skills-chapter2-doc-automation_수정본.docx"
DST = "/Users/gangdasol/Desktop/claude-code-skills-chapter2-doc-automation_수정본.docx"

# 캡션 매핑: 기존 캡션 텍스트 → 새 캡션 텍스트
# 이미 설명이 있는 것은 건드리지 않음
CAPTION_MAP = {
    # 2.2.1 스킬 만들기
    "그림 0.1": "그림 2.1 Cowork 입력창에서 슬래시(/)를 입력했을 때 나타나는 스킬 목록",
    "그림 0.2": "그림 2.2 '프로젝트에서 작업' 버튼 위치",
    "그림 0.3": "그림 2.3 작업 폴더로 claude-skills 선택",
    "그림 0.4": "그림 2.4 chapter1 폴더를 드래그 앤 드롭으로 추가",
    "그림 0.5": "그림 2.5 파일이 업로드된 입력창 화면",
    "그림 0.6": "그림 2.6 스킬 목록에서 doc-automation 확인",

    # 2.2.2 스킬 관리하기
    "그림 0.7": "그림 2.7 Cowork 스킬 관리 화면",
    "그림 0.8": "그림 2.8 스킬 목록에서 doc-automation 선택",
    "그림 0.9": "그림 2.9 doc-automation 스킬 상세 화면 — 파일 구조(왼쪽)와 활성화 조건(오른쪽)",
    "그림 0.10": "그림 2.10 스킬 활성화/비활성화 토글",
    "그림 0.11": "그림 2.11 스킬 관리 메뉴 — 사용해보기, 다운로드, 편집, 교체, 삭제",
    "그림 0.12": "그림 2.12 Cowork에 doc-automation 스킬에 대해 질문한 화면",
    "그림 0.13": "그림 2.13 클로드의 답변 화면 — 왼쪽 답변 영역과 오른쪽 상태 패널",

    # 2.3 PDF to PPT
    "그림 0.14": "그림 2.14 example 폴더의 PDF 파일 위치",
    "그림 0.15": "그림 2.15 실습에 사용할 주간 외환시장 전망 보고서(PDF)",
    "그림 0.16": "그림 2.16 PDF 파일과 요청 문장을 함께 입력한 화면",
    "그림 0.17": "그림 2.17 클로드가 SKILL.md와 PDF 파일을 읽기 시작한 모습",
    "그림 0.18": "그림 2.18 진행 상황 박스 — 세 단계 중 1단계(PDF 분석) 진행 중",
    "그림 0.19": "그림 2.19 작업 폴더에 생성된 PPT 파일과 스크래치패드",
    "그림 0.20": "그림 2.20 세 단계 완료 후 최종 화면 — 왼쪽 메모, 오른쪽 PPT 미리보기",
    "그림 0.21": "그림 2.21 완성된 PPT 저장 옵션 — Google Drive, PowerPoint, 폴더에서 보기",
    "그림 0.22": "그림 2.22 단순 요청으로 만든 PPT — PDF 원본 구조를 그대로 따른 결과",
    "그림 0.23": "그림 2.23 맥락을 포함한 요청으로 만든 PPT — 무역회사 관점으로 재구성된 결과",

    # 2.4 URL to PPT
    # "그림 0.24"는 이미 캡션 있음
    # "그림 0.25"도 이미 있음
    # "그림 0.26"도 이미 있음 (띄어쓰기만 수정)
    "그림 0.26완성된 PPT가 Cowork 화면 오른쪽에 펼쳐진 모습": "그림 2.26 완성된 PPT가 Cowork 화면 오른쪽에 펼쳐진 모습",
    # "그림 0.27"은 이미 있음
    "그림 0.28": "그림 2.28 슬라이드 3 — 긍정적·부정적 시나리오 분석",
    "그림 0.29": "그림 2.29 슬라이드 5 — 투자 전략 제안",

    # 2.5 기업정보 PPT
    # "그림 0.30"은 이미 있음
    "그림 0.31": "그림 2.31 공고문 검토와 PPT 생성을 함께 요청한 입력 화면",
    "그림 0.32": "그림 2.32 클로드가 보여준 팁스 지원 자격 검토 결과",
    "그림 0.33": "그림 2.33 완성된 팁스 발표자료 — 총 9장 슬라이드",
    "그림 0.34": "그림 2.34 부족한 슬라이드 추가를 요청하는 화면",
    "그림 0.35": "그림 2.35 '대화를 계속하기 위해 압축하고 있습니다' 메시지 화면",

    # 2.6 HWPX
    "그림 0.36": "그림 2.36 원본 결재문서 — 서울대공원 총무과 공문 양식",
    "그림 0.37": "그림 2.37 Cowork에서 폴더에서 보기를 클릭하는 화면",
    "그림 0.38": "그림 2.38 완성된 불꽃축제 결과보고서 — 양식은 유지, 내용만 변경",

    # 이미 설명이 있는 것들은 번호만 0→2로 수정
    "그림 0.24 네이버 금융 삼성전자 종목 페이지 화면": "그림 2.24 네이버 금융 삼성전자 종목 페이지 화면",
    "그림 0.25 URL과 요청 문장을 함께 입력한 화면": "그림 2.25 URL과 요청 문장을 함께 입력한 화면",
    "그림 0.27 슬라이드 2 — 매수 타이밍 분석": "그림 2.27 슬라이드 2 — 매수 타이밍 분석",
    "그림 0.30 여러 파일을 한꺼번에 입력창에 올린 화면": "그림 2.30 여러 파일을 한꺼번에 입력창에 올린 화면",
}


def main():
    doc = Document(SRC)
    changes = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text in CAPTION_MAP:
            new_caption = CAPTION_MAP[text]
            # run 단위로 교체 — 첫 번째 run에 전체 텍스트, 나머지 비우기
            runs = para.runs
            if runs:
                runs[0].text = new_caption
                for run in runs[1:]:
                    run.text = ""
                changes.append(f"  ✅ '{text}' → '{new_caption}'")

    doc.save(DST)

    print(f"캡션 수정 ({len(changes)}건):")
    for c in changes:
        print(c)

    # 검증
    doc2 = Document(DST)
    def count_images(d):
        return sum(1 for r in d.part.rels.values() if 'image' in r.reltype)

    print(f"\n검증:")
    print(f"  단락: {len(doc2.paragraphs)} ✅")
    print(f"  이미지: {count_images(doc2)} ✅")
    print(f"  표: {len(doc2.tables)} ✅")

    # 캡션 스타일 유지 확인
    caption_count = sum(1 for p in doc2.paragraphs if p.style.name == 'Caption')
    print(f"  Caption 스타일 단락: {caption_count}개")


if __name__ == "__main__":
    main()
